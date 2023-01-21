# coding: utf8
# from ssl import VERIFY_ALLOW_PROXY_CERTS


from urllib import response
from flask import Flask, render_template, redirect, url_for, jsonify, request, make_response, send_file

# from flask import Flask, flash, jsonify, request, redirect, url_for
from flask_bootstrap import Bootstrap
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, BooleanField
from wtforms.validators import InputRequired, Email, Length
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import create_engine
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from scrape import *
from sentiments import Sentimment_Intensity_Analyzer
from multi_tweets_extraction import *
from visulization import *
import pdfkit
import json
from scrape_youtube import comment_scrape, youtube_info
import subprocess
from transcription import videoDownloader, videoToAudio, audioToTranscript, audioDownload, get_large_audio_transcription, combineChunks, chunksToTranscript, divideChunks
from urduSentiments import test_urdu
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches
import requests
from docx.shared import RGBColor
from docx2pdf import convert
import glob

import shutil
import os
from pathlib import Path
from downloader import download
from spaceModules import m4aToWav, makeChunk

app = Flask(__name__, static_folder='static')
app.config['SECRET_KEY'] = 'Thisissupposedtobesecret!'
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:////home/hassaan/Documents/user_profiling_for_social_media/twitter_app/database.db'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['PROPAGATE_EXCEPTIONS'] = True
bootstrap = Bootstrap(app)
db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

globalDic = {}


class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(15), unique=True)
    email = db.Column(db.String(50), unique=True)
    password = db.Column(db.String(80))


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


class youtubeObject(db.Model):
    vid_id = db.Column(db.BIGINT, primary_key=True)
    query = db.Column(db.String(50))
    Title = db.Column(db.String(100))
    sentiment = db.Column(db.String(10))
    Author = db.Column(db.String(50))
    keywords = db.Column(db.String(2000))
    views = db.Column(db.BIGINT)
    channelId = db.Column(db.String(50))
    link = db.Column(db.String(100))
    publishedOn = db.Column(db.DATETIME)
    thumbnail = db.Column(db.String(100))

    def __init__(self, vid_id, query, Title, sentiment, Author, keywords, views, channelId, link, publishedOn, thumbnail):
        self.vid_id = vid_id
        self.query = query
        self.Title = Title
        self.sentiment = sentiment
        self.Author = Author
        self.keywords = keywords
        self.views = views
        self.channelId = channelId
        self.link = link
        self.publishedOn = publishedOn
        self.thumbnail = thumbnail


@app.before_first_request
def create_table():
    db.create_all()


class LoginForm(FlaskForm):
    print("In login form")
    username = StringField('username', validators=[
                           InputRequired(), Length(min=4, max=15)])
    password = PasswordField('password', validators=[
                             InputRequired(), Length(min=8, max=80)])
    remember = BooleanField('remember me')


class RegisterForm(FlaskForm):
    email = StringField('email', validators=[InputRequired(), Email(
        message='Invalid email'), Length(max=50)])
    username = StringField('username', validators=[
                           InputRequired(), Length(min=4, max=15)])
    password = PasswordField('password', validators=[
                             InputRequired(), Length(min=8, max=80)])


class ScrapeForm(FlaskForm):
    since = StringField('since', validators=[InputRequired()])
    until = StringField('until')
    words = StringField('words')
    hashtag = StringField('hashtag')
    from_account = StringField('from_account')
    mention_account = StringField('mention_account')
    limit = StringField('limit')

    interval = StringField('interval')
    language = StringField('language')
    display_type = StringField('display_type')


####################### Changing from Zarar Side #############################

class YoutubeForm(FlaskForm):
    form = StringField('search', validators=[InputRequired()])

########################Profiles##############################################


@app.route('/profiles', methods=['GET', 'POST'])
def profiles():
    df = pd.read_csv('static/profiles/profiles.csv')
    df.drop('Unnamed: 0', axis=1, inplace=True)
    df['id'] = 0
    for i in range(len(df)):
        df['id'][i] = i
    datadict = df.to_dict('dict')
    return render_template('profiles.html', data=datadict)


@app.route('/profiles/<text>', methods=['GET', 'POST'])
def profileSearcher(text):
    if request.method == 'POST':
        df_dict = {}
        #form = YoutubeForm()
        #text = form.data.get('search')
        print("youtueb query",  text)
        #################Testing Code ########################
        """ df=pd.read_csv('outputs/sentiments_vty.csv')
        df.drop('Unnamed: 0',axis=1,inplace=True)
        df['vid_id']=0
        df['vid_sh_link']=''
        for i in range(len(df)):
            df['vid_id'][i] = i
            df['vid_sh_link'][i] = df['Link'][i].split("v=")[1] """
        #df_dict = df.to_dict('dict')
        # print(df_dict)
        ####################################################
        #### Scraping#######################
        orig = text
        if request.form['submit_button_p'] == 'Slow':
            text = text.replace(" ", "+")
            data_scraped = youtube_info(text)
            df_scraped = pd.DataFrame(data_scraped)
            df_scraped['vid_id'] = 0
            df_scraped['vid_sh_link'] = ''
            df_scraped['query'] = orig
            for i in range(len(df_scraped)):
                df_scraped['vid_id'][i] = i
                df_scraped['vid_sh_link'][i] = df_scraped['Link'][i].split("v=")[
                    1]
            sentiments = Sentimment_Intensity_Analyzer()
            pre_clean_df = sentiments.pre_senti_yt(df_scraped)
            df_vds = pd.DataFrame()
            if pre_clean_df.shape[0] > 0:
                wordcloud_vds_pos, wordcloud_vds_neg, vds_pos, vds_neg, df_vds = sentiments.vds_sentimental(
                    pre_clean_df)
                print(f'Sample of Negative Tweets',
                      df_vds.text_clean[df_vds['sentiments_vds'] == 'neg'])
                print(f'Sample of Postive Tweets ',
                      df_vds.text_clean[df_vds['sentiments_vds'] == 'pos'])
                # 'vid_sh_link','query','Title','sentiments_vds','Author','Keywords','Views',Link,'Published','Thumbnail'
            ## Visulazations ###
                print("Youtube new DF -------->>>", df_vds)
                df_vds.to_csv('outputs/sentiments_vty.csv')
                files = os.listdir('outputs')
                found = False
                for file in files:
                    if 'records' in file:
                        found = True
                        break
                if found == False:
                    df_vds.to_csv('outputs/records.csv')
                else:
                    df_vds.to_csv('outputs/records.csv', mode='a',
                                  index=True, header=False)
                # dict = df.to_dict('dict')
                word_cloud_viz_yt(wordcloud_vds_pos)
                word_cloud_viz_yt(wordcloud_vds_neg,
                                  'WORDCLOUD FOR NEGATIVE YOUTUBE VIDEO', "static/wc_neg_yt.png")
                frequent_words_count_yt(vds_pos)
                frequent_words_count_yt(
                    vds_neg, 'Top words used in Negative Youtube Video', "static/freq_neg_yt.png")
                pubO = list(df_vds['Published_On'])
                print(pubO)
                # print(df_vds)

                df_vds = pd.read_csv('outputs/records.csv')
                df_vds = df_vds[df_vds['query'] == orig]
                df_vds = df_vds.drop_duplicates(
                    subset='vid_sh_link', keep="first")
                if len(df_vds) == 0:
                    return "No data found"

                pubO = list(df_vds['Published_On'])
                print(pubO)
                # print(df_vds)
                pbO = []
                for i in range(len(df_vds)):
                    print("value of i is", i)

                    pbO.append(datetime.strptime(
                        pubO[i], '%Y-%m-%d').strftime('%b %d,%Y'))

                df_vds['pbN'] = pbO
                df_dict = df_vds.to_dict('dict')
                profiles = pd.read_csv('static/profiles/profiles.csv')
                df_dict['query'] = list(
                    profiles[profiles['Name'] == orig]['Name'])[0]
                return render_template('youtubeResults.html', data=df_dict)
        else:
            df_vds = pd.read_csv('outputs/records.csv')
            df_vds = df_vds[df_vds['query'] == orig]
            df_vds = df_vds.drop_duplicates(subset='vid_sh_link', keep="first")
            if len(df_vds) == 0:
                return "No data found"

            pubO = list(df_vds['Published_On'])
            print(pubO)
            # print(df_vds)
            pbO = []
            for i in range(len(df_vds)):
                print("value of i is", i)

                pbO.append(datetime.strptime(
                    pubO[i], '%Y-%m-%d').strftime('%b %d,%Y'))

            df_vds['pbN'] = pbO
            df_dict = df_vds.to_dict('dict')
            profiles = pd.read_csv('static/profiles/profiles.csv')
            df_dict['query'] = list(
                profiles[profiles['Name'] == orig]['Name'])[0]
        # df_scraped.to_csv('outputs/youtube.csv')
        # print("Dataframe of Youtube Information ", df_scraped)
        # link = 'https://www.youtube.com/watch?v=7YQZte-yHm4'
        # data_comments = comment_scrape(link)
        # comments_df = pd.DataFrame(data_comments)
        # print("Comments on Paritcular Video", comments_df)
        # CHANGES FROM ZARAR SIDE
        # df_dict=df_scraped.to_dict('dict')
            return render_template('youtubeResults.html', data=df_dict)
    else:
        return render_template('Wrong Query')
####################### Changing from Zarar Side #############################


@app.route('/pdf/<url>', methods=['GET', 'POST'])
def pdfGen(url):
    dflis = pd.read_csv('outputs/datapdf.csv')
    dflis.drop('Unnamed: 0', axis=1, inplace=True)
    dflis = dflis[dflis['url'] == url]
    lis = dflis.iloc[0].tolist()
    comnts = pd.read_csv('outputs/comments.csv')
    comnts.drop('Unnamed: 0', axis=1, inplace=True)
    # comnts=pd.read_csv('outputs/comments.csv')
    comnts = comnts[comnts['vid_url'] == url]
    response = requests. get(lis[8])
    file = open("static/sample_image.png", "wb")
    file. write(response. content)
    file. close()
    document = Document()
    document.add_heading('Youtube Report', 0)
    document.add_heading('Personal Details', level=2)
    table = document.add_table(rows=5, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    heads = ['Youtuber Name', 'Title', 'Personality Type', 'Watch Url', 'Date']
    vals = [lis[6], lis[5], 'Politician', lis[7], lis[9]]
    titleVideo = "".join(c for c in lis[5] if c.isalpha())
    for i in range(len(heads)):
        data_cells = table.rows[i].cells
        data_cells[0].text = str(heads[i])
        data_cells[1].text = str(vals[i])

    # document.add_heading('Youtuber Name:         '+lis[5],6).style='Normal'
    # document.add_heading('Title:                 '+lis[4],6).style='Normal'
    # document.add_heading('Personality Type:       Politician',6).style='Normal'
    # document.add_heading('Watch Url:             '+lis[6],6).style='Normal'
    # document.add_heading('Date:                  '+lis[8],6).style='Normal'
    document.add_heading('Thumbnails', level=2)
    document.add_picture('static/sample_image.png',
                         width=Inches(5), height=Inches(4)).alignment = 1
    document.add_heading('Sentiments', level=2)
    table = document.add_table(rows=2, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Positive'
    hdr_cells[1].text = 'Negative'
    hdr_cells[2].text = 'Neutral'
    data_cells = table.rows[1].cells
    data_cells[0].text = str(lis[2])
    data_cells[1].text = str(lis[3])
    data_cells[2].text = str(lis[4])
    print(len(lis[0]))
    paragraphs = []
    leng = 600
    paraLen = int(len(lis[1])/leng)
    trans = lis[1]
    for i in range(paraLen+1):
        if i < paraLen:
            paragraphs.append(trans[i*leng:(i*leng)+leng])
        else:
            paragraphs.append(trans[i*leng:])
    print(paragraphs)
    document.add_heading('Transcription', level=2)
    for para in paragraphs:
        p = document.add_paragraph()
        p.alignment = 3
        r = p.add_run()
        font = r.font
        font.complex_script = True
        font.rtl = True
        r.add_text(para)
    # comnts=pd.read_csv('comments.csv')
    # comnts.drop('Unnamed: 0',axis=1,inplace=True)
    document.add_heading('Comments', level=2)
    table = document.add_table(rows=len(comnts)+1, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    com = comnts.to_dict('dict')
    users = list(com['User_Name'].values())
    coments = list(com['Comments'].values())
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'UserName'
    hdr_cells[1].text = 'Comments'
    for i in range(len(users)):
        data_cells = table.rows[i+1].cells
        data_cells[0].text = str(users[i])
        data_cells[1].text = str(coments[i])
    pathToSave = os.path.join("outputs", titleVideo+".docx")
    document.save(pathToSave)

    return send_file(pathToSave, as_attachment=True)


@app.route('/youtube/<url>', methods=['GET', 'POST'])
def youtubeComments(url):
    print(url)
    # found=False
    df = pd.DataFrame()
    if os.path.exists('outputs/datapdf.csv'):
        df = pd.read_csv('outputs/datapdf.csv')
        print('file exist')
    else:
        df = pd.DataFrame(columns=['url', 'trans', 'Neg', 'Neu', 'Pos',
                          'title', 'author', 'Watch_url', 'thumbnail', 'Date'])

    if len(df[df['url'] == url]) == 0 or len(df) == 0:
        updated_url = 'https://www.youtube.com/watch?v='+url
        embed = 'https://www.youtube.com/embed/'+url+'?autoplay=1&mute=0'

        data_comments = comment_scrape(updated_url)
        comments_df = pd.DataFrame(data_comments)
        comments_df['cmnt_id'] = 0
        comments_df['vid_url'] = url
        for i in range(len(comments_df)):
            comments_df['cmnt_id'][i] = i
        print(comments_df)
        print('_______________________')
        print(comments_df.columns)
        if os.path.exists('outputs/comments.csv'):
            comments_df.to_csv('outputs/comments.csv', mode='a', header=False)
        else:
            comments_df.to_csv('outputs/comments.csv')
        """
        comments_df=pd.read_csv('outputs/comments.csv')
        comments_df.drop('Unnamed: 0',axis=1,inplace=True)
        """
        coments_Dict = comments_df.to_dict('dict')
        coments_Dict['EmbedLink'] = embed
        trans = ''
        if videoDownloader(url):
            if videoToAudio(url):
                trans = get_large_audio_transcription(url)
                files = os.listdir("new_ch")
                print(len(files))
                for file in files:
                    os.remove(os.path.join("new_ch", file))
                print(len(os.listdir("new_ch")))
        """ if audioDownload(url):
            subprocess.call(['ffmpeg', '-i', 'audios/test.mp3','audios/'+url+'.wav'])
            trans=audioToTranscript(url) """
        data = pd.read_csv('outputs/records.csv')
        shutil.copy(os.path.join("videos", url+".mp4"),
                    os.path.join("static/playVideo.mp4"))
        desc = list(data[data['vid_sh_link'] == url]['Description'])[0]
        keywords = list(data[data['vid_sh_link'] == url]['Keywords'])[0]
        keyL = keywords.replace('[', '').replace(']', '').replace(
            " '", '').replace("'", '').split(',')
        title = data[data['vid_sh_link'] == url]['Title'].tolist()[0]
        author = data[data['vid_sh_link'] == url]['Author'].tolist()[0]
        thumbnail = data[data['vid_sh_link'] == url]['Thumbnail'].tolist()[0]
        dat = datetime.strptime(data[data['vid_sh_link'] == url]['Published_On'].tolist()[
                                0], '%Y-%m-%d').strftime('%b %d,%Y')
        print(keyL)
        coments_Dict['URL'] = url  # check
        coments_Dict['trans'] = trans
        urduSentiments = test_urdu(trans)
        coments_Dict['Neg'] = "{:.2f}".format(urduSentiments[0]*100)
        coments_Dict['Neu'] = "{:.2f}".format(urduSentiments[1]*100)
        coments_Dict['Pos'] = "{:.2f}".format(urduSentiments[2]*100)
        coments_Dict['desc'] = desc
        coments_Dict['keyList'] = keyL
        coments_Dict['title'] = title
        coments_Dict['author'] = author
        coments_Dict['Watch_url'] = updated_url
        coments_Dict['thumbnail'] = thumbnail
        coments_Dict['Date'] = dat
        datadf = pd.DataFrame(data=[coments_Dict['URL'], coments_Dict['trans'], coments_Dict['Neg'], coments_Dict['Neu'], coments_Dict['Pos'],
                                    coments_Dict['title'], coments_Dict['author'], coments_Dict['Watch_url'], coments_Dict['thumbnail'],
                                    coments_Dict['Date']]).transpose()

        datadf.columns = ['url', 'trans', 'Neg', 'Neu', 'Pos',
                          'title', 'author', 'Watch_url', 'thumbnail', 'Date']
        if os.path.exists('outputs/datapdf.csv'):
            datadf.to_csv('outputs/datapdf.csv', mode='a', header=False)
        else:
            datadf.to_csv('outputs/datapdf.csv')
    else:
        coments = pd.read_csv('outputs/comments.csv')
        coments = coments[coments['vid_url'] == url]
        coments.drop('Unnamed: 0', axis=1, inplace=True)
        coments_Dict = coments.to_dict('dict')
        df = pd.read_csv('outputs/datapdf.csv')
        lisdata = df[df['url'] == url].values.flatten().tolist()
        print("This is our list data")
        print(lisdata)
        # coments_Dict['URL']=lisdata[1]   #check
        coments_Dict['URL'] = url
        ##########################Prininting link of url############################
        print(coments_Dict['URL'])
        coments_Dict['trans'] = lisdata[2]
        #urduSentiments = test_urdu(trans)
        coments_Dict['Neg'] = lisdata[3]
        coments_Dict['Neu'] = lisdata[4]
        coments_Dict['Pos'] = lisdata[5]
        #coments_Dict['desc'] = desc
        #coments_Dict['keyList'] = keyL
        coments_Dict['title'] = lisdata[6]
        coments_Dict['author'] = lisdata[7]
        coments_Dict['Watch_url'] = lisdata[8]
        coments_Dict['thumbnail'] = lisdata[9]
        #datetime.strptime(str(datetime.strptime(df['Timestamp'][0], '%Y-%m-%dT%H:%M:%S.%f%z').date()),'%Y-%m-%d').strftime('%b %d,%Y')
        coments_Dict['Date'] = lisdata[10]
        coments_Dict['EmbedLink'] = 'https://www.youtube.com/embed/' + \
            url+'?autoplay=1&mute=0'
        shutil.copy(os.path.join("videos", url+".mp4"),
                    os.path.join("static/playVideo.mp4"))

    return render_template('youtubeComments.html', data=coments_Dict)


@app.route('/youtube', methods=['GET', 'POST'])
def youtube():

    if request.method == 'POST':
        df_dict = {}
        #form = YoutubeForm()
        text = request.form.get('search')
        #text = form.data.get('search')
        print("youtueb query",  text)
        #################Testing Code ########################

        """ df=pd.read_csv('outputs/sentiments_vty.csv')
        df.drop('Unnamed: 0',axis=1,inplace=True)
        df['vid_id']=0
        df['vid_sh_link']=''
        for i in range(len(df)):
            df['vid_id'][i] = i
            df['vid_sh_link'][i] = df['Link'][i].split("v=")[1] """
        #df_dict = df.to_dict('dict')
        # print(df_dict)

        ####################################################
        #### Scraping#######################
        orig = text
        shTxt = text
        orig = orig.replace(" ", "").lower()
        print("updated youtube query", orig)
        if request.form['submit_button'] == 'Slow':

            text = text.replace(" ", "+")
            data_scraped = youtube_info(text)
            df_scraped = pd.DataFrame(data_scraped)
            df_scraped['vid_id'] = 0
            df_scraped['vid_sh_link'] = ''
            df_scraped['query'] = orig
            for i in range(len(df_scraped)):
                df_scraped['vid_id'][i] = i
                df_scraped['vid_sh_link'][i] = df_scraped['Link'][i].split("v=")[
                    1]

            sentiments = Sentimment_Intensity_Analyzer()
            pre_clean_df = sentiments.pre_senti_yt(df_scraped)
            df_vds = pd.DataFrame()
            if pre_clean_df.shape[0] > 0:
                wordcloud_vds_pos, wordcloud_vds_neg, vds_pos, vds_neg, df_vds = sentiments.vds_sentimental(
                    pre_clean_df)
                print(f'Sample of Negative Tweets',
                      df_vds.text_clean[df_vds['sentiments_vds'] == 'neg'])
                print(f'Sample of Postive Tweets ',
                      df_vds.text_clean[df_vds['sentiments_vds'] == 'pos'])

            ## Visulazations ###
                print("Youtube new DF -------->>>", df_vds)
                df_vds.to_csv('outputs/sentiments_vty.csv')
                files = os.listdir('outputs')
                found = False
                for file in files:
                    if 'records' in file:
                        found = True
                        break
                if found == False:
                    df_vds.to_csv('outputs/records.csv')
                else:
                    df_vds.to_csv('outputs/records.csv', mode='a',
                                  index=True, header=False)
                # dict = df.to_dict('dict')
                word_cloud_viz_yt(wordcloud_vds_pos)
                word_cloud_viz_yt(wordcloud_vds_neg,
                                  'WORDCLOUD FOR NEGATIVE YOUTUBE VIDEO', "static/wc_neg_yt.png")
                frequent_words_count_yt(vds_pos)
                frequent_words_count_yt(
                    vds_neg, 'Top words used in Negative Youtube Video', "static/freq_neg_yt.png")
            df_vds = pd.read_csv('outputs/records.csv')
            df_vds = df_vds[df_vds['query'] == orig]
            df_vds = df_vds.drop_duplicates(subset='vid_sh_link', keep="first")
            if len(df_vds) == 0:
                return "No data found"

            pubO = list(df_vds['Published_On'])
            print(pubO)
            # print(df_vds)
            pbO = []
            for i in range(len(df_vds)):
                print("value of i is", i)

                pbO.append(datetime.strptime(
                    pubO[i], '%Y-%m-%d').strftime('%b %d,%Y'))

            df_vds['pbN'] = pbO
            df_dict = df_vds.to_dict('dict')
            df_dict['query'] = shTxt
        else:
            df_vds = pd.read_csv('outputs/records.csv')
            df_vds = df_vds[df_vds['query'] == orig]
            df_vds = df_vds.drop_duplicates(subset='vid_sh_link', keep="first")
            if len(df_vds) == 0:
                return "No data found"
            pubO = list(df_vds['Published_On'])
            print(pubO)
            # print(df_vds)
            pbO = []
            for i in range(len(df_vds)):
                print("value of i is", i)

                pbO.append(datetime.strptime(
                    pubO[i], '%Y-%m-%d').strftime('%b %d,%Y'))

            df_vds['pbN'] = pbO
            df_dict = df_vds.to_dict('dict')
            df_dict['query'] = shTxt
        # df_scraped.to_csv('outputs/youtube.csv')
        # print("Dataframe of Youtube Information ", df_scraped)
        # link = 'https://www.youtube.com/watch?v=7YQZte-yHm4'
        # data_comments = comment_scrape(link)
        # comments_df = pd.DataFrame(data_comments)
        # print("Comments on Paritcular Video", comments_df)

        # CHANGES FROM ZARAR SIDE

        # df_dict=df_scraped.to_dict('dict')

        return render_template('youtubeResults.html', data=df_dict)
    return render_template('bar.html')


@app.route('/spaceReport', methods=['GET', 'POST'])
def spaceRep():
    dflis = pd.read_csv('outputs/spaceRep.csv')
    dflis.drop('Unnamed: 0', axis=1, inplace=True)
    lis = dflis.iloc[0].tolist()

    document = Document()
    document.add_heading('Twitter Space Report', 0)
    document.add_heading('Personal Details', level=2)
    table = document.add_table(rows=5, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    heads = ['Username', 'Screen Name', 'URL', 'Tweet Time', 'Tweet Date']
    vals = [lis[2], lis[1], lis[0], lis[3], lis[4]]

    for i in range(len(heads)):
        data_cells = table.rows[i].cells
        data_cells[0].text = str(heads[i])
        data_cells[1].text = str(vals[i])

    print("Columns")
    print(dflis.columns)
    print("Columns")
    paragraphs = []
    leng = 600
    paraLen = int(len(lis[7])/leng)
    trans = lis[7]
    for i in range(paraLen+1):
        if i < paraLen:
            paragraphs.append(trans[i*leng:(i*leng)+leng])
        else:
            paragraphs.append(trans[i*leng:])
    print(paragraphs)
    document.add_heading('Transcription', level=2)
    for para in paragraphs:
        p = document.add_paragraph()
        p.alignment = 3
        r = p.add_run()
        font = r.font
        font.complex_script = True
        font.rtl = True
        r.add_text(para)

    document.save('outputs/spaceReport.docx')

    return send_file('outputs/spaceReport.docx', as_attachment=True)


@app.route('/twitterSpace/<url>', methods=['GET', 'POST'])
def spaceshow(url):
    # url='1yoKMZRZqMWGQs'
    try:
        spUrl = 'https://twitter.com/i/spaces/'+url
        if os.path.exists("spaces/"+url+".wav"):
            # divideChunks(url)
            print("Space Exist")
        else:

            # try:
            print('in download space')
            print('chunks combine download space')
            download(url)

            m4aToWav()
            try:
                shutil.move("major.wav", 'spaces/'+url+'.wav')
            except:
                pass
            # divideChunks(url)
        #print('in trans generation space')
        trans = chunksToTranscript(url)

        # trans='sbvfffffffffffffffff'
        print("This is transcript")
        print(trans)
        shutil.copy("spaces/"+url+".wav", "static/space.wav")

        df = pd.read_csv('outputs/spaces.csv')
        dic = {}
        dic['spUrl'] = spUrl
        dic['name'] = df[df['slinks'] == url]['UserScreenName'].tolist()[0]
        dic['uName'] = df[df['slinks'] == url]['UserName'].tolist()[0]
        # dic['ttime']=df[df['slinks']==url]['Timestamp'].tolist()[0]
        dic['ttime'] = df[df['slinks'] == url]['time'].tolist()[0]
        dic['tdate'] = df[df['slinks'] == url]['date'].tolist()[0]
        dic['ttext'] = df[df['slinks'] == url]['Embedded_text'].tolist()[0]
        # dic['likes']= int(df[df['slinks']==url]['Likes'].tolist()[0])
        # dic['retweets']= int(df[df['slinks']==url]['Retweets'].tolist()[0])
        dic['imgLink'] = df[df['slinks'] == url]['Image_link'].tolist()[
            0][2:-2]
        if dic['imgLink'] == '':
            dic['imgLink'] = 'static/space.png'
        dic['splink'] = df[df['slinks'] == url]['splink'].tolist()[0]

        dic['trans'] = trans
        datadf = pd.DataFrame(data=[dic['spUrl'], dic['name'], dic['uName'], dic['ttime'], dic['tdate'],
                                    dic['ttext'], dic['splink'], dic['trans']]).transpose()

        datadf.columns = ['spUrl', 'name', 'uName',
                          'time', 'tdate', 'ttext', 'splink', 'trans']
        datadf.to_csv('outputs/spaceRep.csv')

        return render_template('space.html', data=dic)
    except:
        return "Oops Some thing bad happend"


@app.route('/twitterSpace', methods=['GET', 'POST'])
def twitterSpace():
    if request.method == 'POST':
        # https://twitter.com/i/spaces/1yoKMZRZqMWGQ?s=20
        since = request.form.get('since')
        since = str(since)
        until = request.form.get('until')
        until = str(until)
        language = request.form.get('language')
        language = str(language)

        word = request.form.get('words')
        word = str(word)
        wor = 'filter:spaces'
        retrived_data = scrape(words=wor, until=until, since=since,
                               limit=10, interval=1, display_type='Latest', lang=language, headless=True)

        df = retrived_data
        print(df)
        df = pre_clean_df_spaces(df)
        print('check point 0')
        df['tweet_id'] = 0
        for i in range(len(df)):
            df['tweet_id'][i] = i+1
        print('check point 1')
        df.rename(columns={'Image link': 'Image_link'}, inplace=True)
        df.to_csv('outputs/spaces.csv')
        df = pd.read_csv('outputs/spaces.csv')
        df.drop('Unnamed: 0', axis=1, inplace=True)

        print(df)
        df['slinks'] = ''
        df['time'] = ''
        df['date'] = ''
        for i in range(len(df)):
            try:
                df['slinks'][i] = df['splink'][i].split('spaces/')[1]
            except:
                pass
            df['time'][i] = datetime.strptime(
                df['Timestamp'][i], '%Y-%m-%dT%H:%M:%S.%f%z').time()
            df['date'][i] = datetime.strptime(str(datetime.strptime(
                df['Timestamp'][0], '%Y-%m-%dT%H:%M:%S.%f%z').date()), '%Y-%m-%d').strftime('%b %d,%Y')

        df.to_csv('outputs/spaces.csv')
        df = pd.read_csv('outputs/spaces.csv')

        df.dropna(subset=['UserScreenName'], inplace=True)

        dict = df.to_dict('dict')

        return render_template("twitterSpaceResults.html", data=dict)

    return render_template('twitterspace.html')

######################## End Changing from Zarar Side ########################


@app.route('/index', methods=['GET', 'POST'])
def index():
    return render_template('index.html')


@app.route('/analytics')
def analytics():
    return render_template('analytics.html')


@app.route('/', methods=['GET', 'POST'])
def login():
    print('In login')
    form = LoginForm()

    if form.validate_on_submit():
        print('validate_on_submit()')
        user = User.query.filter_by(username=form.username.data).first()
        if user:
            print('in user')
            if check_password_hash(user.password, form.password.data):
                print('in check password hash')
                login_user(user, remember=form.remember.data)
                return redirect(url_for('index'))
            else:
                print("Invalid password")
        return '<h1>Invalid username or password</h1>'

    return render_template('login.html', form=form)


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    form = RegisterForm()

    if form.validate_on_submit():
        hashed_password = generate_password_hash(
            form.password.data, method='sha256')
        new_user = User(username=form.username.data,
                        email=form.email.data, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()

        return redirect(url_for('login'))

    return render_template('signup.html', form=form)


@app.route('/tweet_scrape', methods=['GET', 'POST'])
def scaper():

    form = ScrapeForm()
    since = form.data.get('since')
    since = str(since)
    until = form.data.get('until')
    until = str(until)
    words = form.data.get('words')
    hashtag = form.data.get('hashtag')
    from_account = form.data.get('from_account')
    from_account = str(from_account)
    mention_account = form.data.get('mention_account')
    mention_account = str(mention_account)
    limit = form.data.get('limit')
    limit = int(limit)
    interval = form.data.get('interval')
    interval = int(interval)
    language = form.data.get('language')
    language = str(language)
    display_type = form.data.get('display_type')
    display_type = str(display_type)

    ############################### JUST FOR TESTING ##################################
    """ df=pd.read_csv('outputs/hashtag_tweets_new.csv')
    dict=df.to_dict()
    return render_template('results.html',data=dict) """
    ############################## JUST FOR TESTING####################################
    tweets = scrape(hashtag=hashtag, words=words, since=since, until=until, from_account=from_account,
                    to_account=None, mention_account=mention_account, interval=interval, headless=True, display_type=display_type, save_images=False,
                    proxy=None, save_dir='outputs', lang=language, resume=False, filter_replies=True, proximity=False,  limit=limit,
                    show_images=False,  geocode=None, minreplies=None, minlikes=None, minretweets=None)

    mlt = Multi_Tweets_Extraction()
    sentiments = Sentimment_Intensity_Analyzer()
    mlt_df = mlt.multi_tweets(tweets)
    df = pd.DataFrame(mlt_df)
    pre_clean_df = sentiments.pre_senti(df)

    if pre_clean_df.shape[0] > 0:
        wordcloud_vds_pos, wordcloud_vds_neg, vds_pos, vds_neg, df_vds = sentiments.vds_sentimental(
            pre_clean_df)
        print(f'Sample of Negative Tweets',
              df_vds.text_clean[df_vds['sentiments_vds'] == 'neg'])
        print(f'Sample of Postive Tweets ',
              df_vds.text_clean[df_vds['sentiments_vds'] == 'pos'])

    ## Visulazations ###
        print(df_vds)
        df_vds.to_csv('twitterData.csv')
        dict = df.to_dict('dict')
        word_cloud_viz(wordcloud_vds_pos)
        word_cloud_viz(wordcloud_vds_neg,
                       'WORDCLOUD FOR NEGATIVE TWEETS', "static/wc_neg.png")
        frequent_words_count(vds_pos)
        frequent_words_count(
            vds_neg, 'Top words used in Negative Tweets', "static/freq_neg.png")
        return render_template('results.html', data=dict)
    else:
        return "no data found"


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', name=current_user.username)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


if __name__ == '__main__':
    app.run(debug=True)
